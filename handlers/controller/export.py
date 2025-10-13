from aiogram import Router, F
from aiogram.types import Message, CallbackQuery, BufferedInputFile
from aiogram.fsm.context import FSMContext
from filters.role_filter import RoleFilter
from keyboards.controllers_buttons import get_controller_export_types_keyboard, get_controller_export_formats_keyboard
from database.controller.export import (
    get_controller_orders_for_export,
    get_controller_statistics_for_export,
    get_controller_employees_for_export,
)
from utils.export_utils import ExportUtils
from utils.universal_error_logger import get_universal_logger, log_error
import logging
from datetime import datetime

router = Router()
router.message.filter(RoleFilter(role="controller"))
logger = get_universal_logger("ControllerExport")

@router.message(F.text.in_(["📤 Export", "📤 Экспорт"]))
async def export_handler(message: Message, state: FSMContext):
    """Main export handler - shows export types"""
    try:
        await state.clear()
        keyboard = get_controller_export_types_keyboard()
        await message.answer(
            "📊 <b>Kontrollerlar uchun hisobotlar</b>\n\n"
            "Quyidagi hisobot turlaridan birini tanlang:",
            reply_markup=keyboard,
            parse_mode="HTML"
        )
    except Exception as e:
        logger.error(f"Export handler error: {e}")
        await message.answer("❌ Xatolik yuz berdi. Iltimos, qaytadan urinib ko'ring.")

@router.callback_query(F.data == "controller_export_tech_requests")
async def export_tech_requests_handler(callback: CallbackQuery, state: FSMContext):
    """Handle tech requests export selection"""
    try:
        await state.update_data(export_type="tech_requests")
        keyboard = get_controller_export_formats_keyboard()
        await callback.message.edit_text(
            "📋 <b>Texnik arizalar ro'yxati</b>\n\n"
            "Export formatini tanlang:",
            reply_markup=keyboard,
            parse_mode="HTML"
        )
        await callback.answer()
    except Exception as e:
        logger.error(f"Export tech requests handler error: {e}")
        await callback.answer("❌ Xatolik yuz berdi", show_alert=True)



@router.callback_query(F.data == "controller_export_statistics")
async def export_statistics_handler(callback: CallbackQuery, state: FSMContext):
    """Handle statistics export selection"""
    try:
        await state.update_data(export_type="statistics")
        keyboard = get_controller_export_formats_keyboard()
        await callback.message.edit_text(
            "📊 <b>Statistika hisoboti</b>\n\n"
            "Export formatini tanlang:",
            reply_markup=keyboard,
            parse_mode="HTML"
        )
        await callback.answer()
    except Exception as e:
        logger.error(f"Export statistics handler error: {e}")
        await callback.answer("❌ Xatolik yuz berdi", show_alert=True)

@router.callback_query(F.data == "controller_export_employees")
async def export_employees_handler(callback: CallbackQuery, state: FSMContext):
    """Handle employees export selection"""
    try:
        await state.update_data(export_type="employees")
        keyboard = get_controller_export_formats_keyboard()
        await callback.message.edit_text(
            "👥 <b>Xodimlar ro'yxati</b>\n\n"
            "Export formatini tanlang:",
            reply_markup=keyboard,
            parse_mode="HTML"
        )
        await callback.answer()
    except Exception as e:
        logger.error(f"Export employees handler error: {e}")
        await callback.answer("❌ Xatolik yuz berdi", show_alert=True)


@router.callback_query(F.data.startswith("controller_format_"))
async def export_format_handler(callback: CallbackQuery, state: FSMContext):
    """Handle export format selection and generate file"""
    try:
        format_type = callback.data.split("_")[-1]  # csv, xlsx, docx, pdf
        data = await state.get_data()
        export_type = data.get("export_type", "tech_requests")
        
        # Show processing message
        await callback.message.edit_text(
            "⏳ <b>Hisobot tayyorlanmoqda...</b>\n\n"
            "Iltimos, kuting...",
            parse_mode="HTML"
        )
        
        # Get data based on export type
        if export_type == "tech_requests":
            orders_data = await get_controller_orders_for_export()
            title = "Texnik arizalar ro'yxati"
            filename_base = "texnik_arizalar"
            headers = [
                "ID", "Ariza raqami", "Mijoz ismi", "Telefon",
                "Manzil", "Ish tavsifi", "Holati",
                "Texnik", "Kontroller",
                "Yaratilgan sana", "Yangilangan sana",
                "Akt raqami"
            ]
            
            # Convert dict data to list format for export
            raw_data = [
                [
                    order.get("id", ""),
                    order.get("application_number", ""),
                    order.get("client_name", ""),
                    order.get("client_phone", ""),
                    order.get("address", ""),
                    order.get("description", ""),
                    order.get("status", ""),
                    order.get("technician_name", ""),  # Texnik
                    order.get("controller_name", ""),  # Kontroller
                    order.get("created_at", ""),
                    "",  # Yangilangan sana (bo'sh)
                    order.get("akt_number", "")  # Akt raqami
                ]
                for order in orders_data
            ]
            
        elif export_type == "connection_orders":
            raw_data = await get_controller_orders_for_export()
            title = "Ulanish arizalari ro'yxati"
            filename_base = "ulanish_arizalari"
            headers = [
                "ID", "Ariza raqami", "Mijoz ismi", "Telefon",
                "Manzil", "Tarif rejasi", "Holati",
                "Ulanish sanasi", "Yangilangan sana",
                "Akt raqami"
            ]
            
        elif export_type == "statistics":
            stats = await get_controller_statistics_for_export()

            if not stats or 'summary' not in stats:
                logger.error("Failed to get statistics for export or summary is missing.")
                await callback.message.answer(
                    "❌ Statistika ma'lumotlarini olishda xatolik yuz berdi.\n"
                    "Iltimos, keyinroq qayta urinib ko'ring."
                )
                await callback.answer()
                return
            raw_data = []
            title = "Statistika hisoboti"
            filename_base = "statistika"
            
            def add_section(title):
                nonlocal raw_data
                raw_data.append(["", ""])
                raw_data.append([f"🔹 {title.upper()}", ""])
                raw_data.append(["-" * 30, "-" * 30])
            
            def add_row(label, value, indent=0):
                nonlocal raw_data
                prefix = "  " * indent
                raw_data.append([f"{prefix}{label}", str(value) if value is not None else "0"])
            
            # 1. Asosiy statistika
            add_section("Umumiy statistika")
            add_row("📊 Jami texnik arizalar:", stats['summary']['total_requests'])
            add_row("🆕 Yangi arizalar:", stats['summary']['new_requests'])
            add_row("🔄 Jarayondagi arizalar:", stats['summary']['in_progress_requests'])
            add_row("✅ Yakunlangan arizalar:", stats['summary']['completed_requests'])
            add_row("📈 Yakunlangan arizalar foizi:", f"{stats['summary'].get('completion_rate', 0)}%")
            add_row("👥 Yagona mijozlar:", stats['summary']['unique_clients'])
            add_row("🔧 Muammo turlari:", stats['summary'].get('unique_tariffs', 0))
            
            # 2. Texniklar bo'yicha statistika
            if stats['by_technician']:
                add_section("Texniklar bo'yicha statistika")
                for i, technician in enumerate(stats['by_technician'], 1):
                    technician_name = f"👤 {i}. {technician['technician_name']}"
                    phone = technician['technician_phone'] or 'Tel. yo\'q'
                    add_row(technician_name, "", 0)
                    add_row("  📞 Telefon:", phone, 1)
                    add_row("  📊 Jami arizalar:", technician['total_orders'], 1)
                    add_row("  ✅ Yakunlangan:", technician['completed_orders'], 1)
                    add_row("  🔄 Jarayonda:", technician['in_progress_orders'], 1)
                    raw_data.append(["", ""])  # Empty row after each technician
            
            # 3. Oylik statistika
            if stats.get('monthly_trends'):
                add_section("Oylik statistika (6 oy)")
                for month_data in stats['monthly_trends']:
                    month = month_data['month']
                    add_row(f"🗓️ {month}:", "", 0)
                    add_row("  📊 Jami:", month_data['total_requests'], 1)
                    add_row("  🆕 Yangi:", month_data['new_requests'], 1)
                    add_row("  ✅ Yakunlangan:", month_data['completed_requests'], 1)
            
            # 4. Muammo turlari bo'yicha statistika (currently not implemented)
            # if stats['by_problem_type']:
            #     add_section("Muammo turlari bo'yicha statistika")
            #     for problem in stats['by_problem_type']:
            #         add_row(f"🔧 {problem['problem_type']}", "", 0)
            #         add_row("  📊 Arizalar soni:", problem['total_requests'], 1)
            #         add_row("  👥 Mijozlar soni:", problem['unique_clients'], 1)
            #         add_row("  ✅ Yakunlangan:", problem['completed_requests'], 1)
            
            # 5. So'nggi faollik
            if stats['recent_activity']:
                add_section("So'nggi faollik (30 kun)")
                for activity in stats['recent_activity']:
                    if activity['recent_orders'] > 0:
                        last_active = activity['last_order_date'].strftime('%Y-%m-%d') if activity['last_order_date'] else 'Noma\'lum'
                        add_row(
                            f"👤 {activity['technician_name']}",
                            f"📅 So'nggi: {last_active}",
                            0
                        )
                        add_row("  📊 Arizalar soni:", activity['recent_orders'], 1)
                
            headers = ["Ko'rsatkich", "Qiymat"]
            
        elif export_type == "employees":
            employees = await get_controller_employees_for_export()
            title = "Xodimlar ro'yxati"
            filename_base = "xodimlar"
            headers = [
                "Ism-sharif", "Telefon", "Lavozim",
                "Qo'shilgan sana"
            ]
            
            # Convert the list of dicts to the format expected by the export functions
            raw_data = [
                [
                    emp.get("full_name", ""),
                    emp.get("phone", ""),
                    emp.get("role", ""),
                    emp.get("created_at", "").strftime('%Y-%m-%d') if emp.get("created_at") else ""
                ]
                for emp in employees
            ]
            
        elif export_type == "reports":
            raw_data = await get_controller_reports_for_export()
            title = "Hisobotlar"
            filename_base = "hisobotlar"
            headers = [
                "Sarlavha", "Yaratuvchi", 
                "Holati", "Yaratilgan sana"
            ]
        
        # Generate file based on format
        try:
            if format_type == "xlsx":
                file = await generate_excel(raw_data, headers, title, filename_base)
            elif format_type == "csv":
                file = await generate_csv(raw_data, headers, title, filename_base)
            elif format_type == "docx":
                # For Word export, ensure data is in the correct format
                if export_type in ["employees", "reports"]:
                    # For these types, raw_data is already a list of dicts
                    file = await generate_word(raw_data, headers, title, filename_base)
                else:
                    # For other types, convert to list of dicts
                    dict_data = _rows_to_dicts(raw_data, headers)
                    file = await generate_word(dict_data, headers, title, filename_base)
            elif format_type == "pdf":
                file = await generate_pdf(raw_data, headers, title, filename_base)
            else:
                raise ValueError("Noto'g'ri format tanlandi")
        except Exception as e:
            logger.error(f"Error generating {format_type.upper()} file: {str(e)}", exc_info=True)
            raise ValueError(f"{format_type.upper()} faylini yaratishda xatolik: {str(e)}")
        
        # Send the file
        await callback.message.answer_document(
            document=file,
            caption=f"📤 {title}\n"
                   f"📅 {datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
                   f"✅ Muvaffaqiyatli yuklab olindi!"
        )
        
        await state.clear()
        
    except Exception as e:
        log_error(e, "Controller export format handler", callback.from_user.id)
        await callback.message.answer(
            f"❌ Xatolik yuz berdi: {str(e)}\n"
            "Iltimos, qaytadan urinib ko'ring yoki administratorga murojaat qiling."
        )
    
    await callback.answer()

@router.callback_query(F.data == "controller_export_back_types")
async def export_back_to_types_handler(callback: CallbackQuery, state: FSMContext):
    """Handle back to export types"""
    try:
        keyboard = get_controller_export_types_keyboard()
        await callback.message.edit_text(
            "📊 <b>Kontrollerlar uchun hisobotlar</b>\n\n"
            "Quyidagi hisobot turlaridan birini tanlang:",
            reply_markup=keyboard,
            parse_mode="HTML"
        )
        await callback.answer()
    except Exception as e:
        logger.error(f"Export back to types handler error: {e}")
        await callback.answer("❌ Xatolik yuz berdi", show_alert=True)

@router.callback_query(F.data == "controller_export_end")
async def export_end_handler(callback: CallbackQuery, state: FSMContext):
    """Handle export end"""
    try:
        await callback.message.delete()
        await state.clear()
        await callback.answer()
    except Exception as e:
        logger.error(f"Export end handler error: {e}")
        await callback.answer("❌ Xatolik yuz berdi", show_alert=True)

def _rows_to_dicts(data: list, headers: list) -> list[dict]:
    """Convert rows to list of dicts with only the specified headers"""
    dict_data = []
    if not data:
        return dict_data

    # If rows are already dicts, just filter the keys
    if isinstance(data[0], dict):
        for row in data:
            row_dict = {}
            for header in headers:
                # Try direct key access first, then try with header mapping
                value = row.get(header) or row.get(_get_db_key_for_header(header), "")
                row_dict[header] = str(value) if value is not None else ""
            dict_data.append(row_dict)
        return dict_data

    # Handle list of lists/tuples
    for row in data:
        if not isinstance(row, (list, tuple, dict)):
            # Single value case
            row_dict = {headers[0]: str(row) if row is not None else ""}
            # Add empty values for remaining headers
            for header in headers[1:]:
                row_dict[header] = ""
        else:
            # List/tuple case
            row_dict = {}
            for i, header in enumerate(headers):
                if i < len(row):
                    value = row[i]
                    row_dict[header] = str(value) if value is not None else ""
                else:
                    row_dict[header] = ""
        
        dict_data.append(row_dict)
    
    return dict_data


def _get_db_key_for_header(header: str) -> str:
    mapping = {
        # Common fields
        "ID": "id",
        "Holati": "status",
        "Yaratilgan sana": "created_at",
        "Yangilangan sana": "updated_at",
        
        # Client/technician/controller
        "Ariza raqami": "request_number",
        "Mijoz ismi": "client_name",
        "Telefon": "phone_number",
        "Mijoz abonent ID": "client_abonent_id",
        "Texnik": "assigned_technician",
        "Kontroller": "controller_name",
        
        # Address/work/tariff
        "Hudud": "region",
        "Manzil": "address",
        "Ish tavsifi": "description_ish",
        "Tarif rejasi": "plan_name",
        "Ulanish sanasi": "connection_date",
        
        # AKT documents
        "Akt raqami": "akt_number",
        "Akt fayl yo'li": "akt_file_path",
        "Akt yaratilgan": "akt_created_at",
        "Mijozga yuborilgan": "sent_to_client_at",
        "Akt reytingi": "akt_rating",
        "Akt izohi": "akt_comment",
        
        # Additional fields
        "Abonent ID": "abonent_id",
        "Media": "media",
        "Uzunlik": "longitude",
        "Kenglik": "latitude",
        "Muammo turi": "description",
        "Tavsif": "description",
        "Reyting": "rating",
        "Izohlar": "notes",
        "Texnik telefon": "technician_phone",
        "Kontroller telefon": "controller_phone"
    }
    return mapping.get(header, header.lower().replace(" ", "_")).replace(" ", "_")

async def generate_excel(data: list, headers: list, title: str, filename: str) -> BufferedInputFile:
    """Generate Excel file from data"""
    dict_data = _rows_to_dicts(data, headers)
    if not data:
        # Handle empty data
        dict_data = []
    elif isinstance(data[0], dict):
        # Data is already in dictionary format (from database queries)
        for row in data:
            row_dict = {}
            for header in headers:
                # Map header to corresponding database column key
                 row_dict[header] = row.get(_get_db_key_for_header(header), "")
            dict_data.append(row_dict)
    else:
        # Data is in list format (like statistics)
        for row in data:
            row_dict = {}
            for i, header in enumerate(headers):
                if i < len(row):
                    row_dict[header] = row[i]
                else:
                    row_dict[header] = ""
            dict_data.append(row_dict)
    
    # Use ExportUtils to generate Excel
    output = ExportUtils.generate_excel(dict_data, title[:30], title)
    
    return BufferedInputFile(
        file=output.getvalue(),
        filename=f"{filename}_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
    )

async def generate_csv(data: list, headers: list, title: str, filename: str) -> BufferedInputFile:
    """Generate CSV file from data"""
    import io
    import csv
    
    # Convert data to list of dicts with only the specified headers
    dict_data = _rows_to_dicts(data, headers)
    
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=headers)
    
    # Write headers
    writer.writeheader()
    
    # Write data rows
    for row in dict_data:
        writer.writerow(row)
    
    return BufferedInputFile(
        file=output.getvalue(),
        filename=f"{filename}_{datetime.now().strftime('%Y%m%d_%H%M')}.csv"
    )

async def generate_word(data: list, headers: list, title: str, filename: str) -> BufferedInputFile:
    """Generate Word file from data"""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io
    
    # Convert data to list of dicts with only the specified headers
    dict_data = _rows_to_dicts(data, headers)
    
    doc = Document()
    
    # Add title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Sana: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    date_run.italic = True
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Add headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
    
    # Add data rows
    for row in dict_data:
        row_cells = table.add_row().cells
        for i, header in enumerate(headers):
            row_cells[i].text = str(row.get(header, ""))
    
    # Save to bytes
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    return BufferedInputFile(
        file=output.read(),
        filename=f"{filename}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    )

async def generate_pdf(data: list, headers: list, title: str, filename: str) -> BufferedInputFile:
    """Generate PDF file from data"""
    # Convert data to list of dicts with only the specified headers
    dict_data = _rows_to_dicts(data, headers)
    
    # Use ExportUtils to generate PDF
    output = ExportUtils.generate_pdf(dict_data, title)
    
    return BufferedInputFile(
        file=output.getvalue(),
        filename=f"{filename}_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf"
    )
